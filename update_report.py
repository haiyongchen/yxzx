# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print('小橘开始更新报告...')

# 读取所有数据
df_c = pd.read_csv('D:\\openclaw-workspace\\all_zones_c_contract.csv', encoding='utf-8-sig')
df_2026 = pd.read_csv('D:\\openclaw-workspace\\zones_2026.csv', encoding='utf-8-sig')

# 数据处理 - 修正：专区成本已包含上线前成本
df_c['总收益情况'] = pd.to_numeric(df_c['总收益情况'], errors='coerce').fillna(0)
df_c['上线前成本'] = pd.to_numeric(df_c['上线前成本'], errors='coerce').fillna(0)
df_c['专区成本'] = pd.to_numeric(df_c['专区成本'], errors='coerce').fillna(0)
# 关键修正：使用专区成本作为总成本（已包含上线前成本）
df_c['总成本'] = df_c['专区成本']

# 处理时间
df_c['确认接入时间'] = pd.to_datetime(df_c['确认接入时间'], errors='coerce')
df_c['接入年份'] = df_c['确认接入时间'].dt.year

# 创建文档
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

# 标题
title = doc.add_heading('e交易专区收益成本统计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('（合同编号C开头 - 110个专区维度分析）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
doc.add_paragraph('中原区、华北区运营分析报告').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('报告日期：2026年3月31日').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)
doc.add_paragraph('截至2026年3月31日，新点e交易（中原、华北区）合同编号以C开头的专区共运营110个，覆盖10个省份，累计收益4,145,231.71元，累计成本1,488,500.37元。')
doc.add_paragraph('【重要说明】专区成本已包含上线前成本，不再重复计算。')

# 关键指标
doc.add_heading('【关键指标】', 2)
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['指标名称', '数值', '说明']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

total_cost = df_c['专区成本'].sum()
avg_cost = df_c['专区成本'].mean()
high_cost_count = len(df_c[df_c['专区成本'] > 32000])

data = [
    ['专区总数', '110个', '合同编号C开头'],
    ['覆盖省份', '10个', '湖北、新疆、河南等'],
    ['累计收益', '4,145,231.71元', '-'],
    ['累计成本', f'{total_cost:,.2f}元', '专区成本（含上线前成本）'],
    ['平均单专区收益', '37,683.92元', '-'],
    ['平均单专区成本', f'{avg_cost:,.2f}元', '专区成本（含上线前成本）'],
    ['成本超标专区', f'{high_cost_count}个', '专区成本>32,000元'],
    ['26年度新开设', '8个', '2026年接入'],
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

print('关键指标完成')

# 主要问题
doc.add_heading('【主要问题】', 2)
problems = [
    '河北、内蒙古、辽宁、吉林等省份成本率严重超标（>100%）',
    f'{high_cost_count}个专区专区成本超过32,000元基线',
    '部分专区零收益，成本无法回收',
    '26年度新开设8个专区目前收益均为0，需关注后续上量',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# 二、核心数据总览
doc.add_heading('二、核心数据总览', 1)

# 2.1 各省份专区分布
doc.add_heading('2.1 各省份专区分布（修正后）', 2)
province_stats = df_c.groupby('所属省份').agg({
    '合同编号': 'count',
    '总收益情况': 'sum',
    '专区成本': 'sum'
}).reset_index()
province_stats.columns = ['省份', '专区数量', '累计收益', '累计成本']
province_stats['成本率'] = province_stats.apply(lambda x: f"{(x['累计成本'] / x['累计收益'] * 100):.1f}%" if x['累计收益'] > 0 else 'N/A', axis=1)
province_stats = province_stats.sort_values('累计收益', ascending=False)

table2 = doc.add_table(rows=len(province_stats)+1, cols=5)
table2.style = 'Light Grid Accent 1'
headers2 = ['省份', '专区数量', '累计收益', '累计成本(专区成本)', '成本率']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, row in enumerate(province_stats.itertuples(), 1):
    table2.rows[idx].cells[0].text = str(row.省份)
    table2.rows[idx].cells[1].text = str(int(row.专区数量))
    table2.rows[idx].cells[2].text = f"{row.累计收益:,.0f}"
    table2.rows[idx].cells[3].text = f"{row.累计成本:,.0f}"
    table2.rows[idx].cells[4].text = str(row.成本率)

print('省份分布完成')

doc.add_page_break()

# 三、26年度新开设专区专项分析（新增）
doc.add_heading('三、26年度新开设专区专项分析', 1)
doc.add_paragraph('本章节针对2026年度新开设的8个专区进行专项成本分析，以评估新专区成本控制情况。')

# 3.1 总体情况
doc.add_heading('3.1 总体情况', 2)
table_2026_1 = doc.add_table(rows=7, cols=2)
table_2026_1.style = 'Light Grid Accent 1'
table_2026_1.rows[0].cells[0].text = '指标'
table_2026_1.rows[0].cells[1].text = '数值'
table_2026_1.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
table_2026_1.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

stats_2026 = [
    ['新开设专区数', '8个'],
    ['总成本', f"{df_2026['专区成本'].sum():,.2f}元"],
    ['平均成本', f"{df_2026['专区成本'].mean():,.2f}元"],
    ['成本中位数', f"{df_2026['专区成本'].median():,.2f}元"],
    ['最高成本', f"{df_2026['专区成本'].max():,.2f}元"],
    ['最低成本', f"{df_2026['专区成本'].min():,.2f}元"],
]
for i, row in enumerate(stats_2026, 1):
    table_2026_1.rows[i].cells[0].text = row[0]
    table_2026_1.rows[i].cells[1].text = row[1]

# 3.2 成本分布
doc.add_heading('3.2 成本分布分析', 2)
high = df_2026[df_2026['专区成本'] > 32000]
med = df_2026[(df_2026['专区成本'] > 10000) & (df_2026['专区成本'] <= 32000)]
low = df_2026[df_2026['专区成本'] <= 10000]

doc.add_paragraph(f'高成本专区(>32,000元)：{len(high)}个')
doc.add_paragraph(f'中成本专区(10,000-32,000元)：{len(med)}个')
doc.add_paragraph(f'低成本专区(≤10,000元)：{len(low)}个')
doc.add_paragraph()
doc.add_paragraph('【分析结论】26年度新开设专区成本控制良好，无超标专区（>32,000元），平均成本9,500元远低于基线。')

# 3.3 专区明细
doc.add_heading('3.3 26年度新开设专区明细', 2)
table_2026_2 = doc.add_table(rows=len(df_2026)+1, cols=5)
table_2026_2.style = 'Light Grid Accent 1'
headers_2026 = ['合同编号', '专区名称', '省份', '接入时间', '专区成本']
for i, header in enumerate(headers_2026):
    table_2026_2.rows[0].cells[i].text = header
    table_2026_2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(df_2026.iterrows(), 1):
    table_2026_2.rows[idx].cells[0].text = str(row['合同编号'])
    table_2026_2.rows[idx].cells[1].text = str(row[' 原专区名称'])[:15]
    table_2026_2.rows[idx].cells[2].text = str(row['所属省份'])
    table_2026_2.rows[idx].cells[3].text = str(row['确认接入时间'])[:10]
    table_2026_2.rows[idx].cells[4].text = f"{row['专区成本']:,.2f}"

print('26年度分析完成')

doc.add_page_break()

# 四、数据分析与洞察
doc.add_heading('四、数据分析与洞察', 1)

# 4.1 成本分析
doc.add_heading('4.1 成本分析（修正后）', 2)
doc.add_paragraph('【重要说明】以下成本分析基于"专区成本"字段，该字段已包含上线前成本，不再重复计算。')
cost_analysis = [
    f'平均专区成本：{avg_cost:,.2f}元（远低于基线32,000元）',
    f'成本超标专区：{high_cost_count}个（专区成本>32,000元）',
    '成本率超标省份：河北(347%)、内蒙古(180%)、辽宁(132%)、吉林(1761%)',
    '整体成本率：35.9%（健康水平）',
    '26年度新开设专区平均成本：9,500元，无超标',
]
for item in cost_analysis:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 五、重点工作规划
doc.add_heading('五、重点工作规划', 1)

# 5.1 26年度新专区跟进
doc.add_heading('5.1 26年度新开设专区跟进计划', 2)
doc.add_paragraph('责任人：XXX    完成时间：2026年6月30日')
doc.add_paragraph('跟进措施：')
plans_2026 = [
    '重点关注中成本专区（4个），确保上量进度',
    '定期跟踪收益情况，目标Q2实现首笔收益',
    '对零收益专区进行专项分析，制定上量方案',
]
for p in plans_2026:
    doc.add_paragraph(p, style='List Bullet')

# 保存
output_path = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\e交易专区收益成本统计报告_C合同_110个专区_修正版.docx'
doc.save(output_path)

print(f'报告更新完成！')
print(f'文件路径: {output_path}')
print('主要更新:')
print('  1. 修正成本计算：使用专区成本（已含上线前成本）')
print('  2. 新增26年度新开设专区分析章节')
print('  3. 更新所有相关表格和指标')
