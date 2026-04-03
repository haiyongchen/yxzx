# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 读取数据
df_c = pd.read_csv('D:\\openclaw-workspace\\all_zones_c_contract.csv', encoding='utf-8-sig')

# 数据处理
df_c['总收益情况'] = pd.to_numeric(df_c['总收益情况'], errors='coerce').fillna(0)
df_c['上线前成本'] = pd.to_numeric(df_c['上线前成本'], errors='coerce').fillna(0)
df_c['专区成本'] = pd.to_numeric(df_c['专区成本'], errors='coerce').fillna(0)
df_c['总成本'] = df_c['上线前成本'] + df_c['专区成本']
df_c['成本率'] = df_c.apply(lambda x: (x['总成本'] / x['总收益情况'] * 100) if x['总收益情况'] > 0 else 999999, axis=1)

# 创建文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

# 标题
title = doc.add_heading('e交易专区收益成本统计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph('（合同编号C开头 - 110个专区维度分析）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)

# 报告信息
doc.add_paragraph('中原区、华北区运营分析报告').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('报告日期：2026年3月31日').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)
doc.add_paragraph('截至2026年3月31日，新点e交易（中原、华北区）合同编号以C开头的专区共运营110个，覆盖10个省份，累计收益4,145,231.71元，累计成本1,488,500.37元，平均单专区收益37,683.92元。')

# 关键指标
doc.add_heading('【关键指标】', 2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['指标名称', '数值', '说明']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    ['专区总数', '110个', '合同编号C开头'],
    ['覆盖省份', '10个', '湖北、新疆、河南等'],
    ['累计收益', '4,145,231.71元', '-'],
    ['累计成本', '1,488,500.37元', '-'],
    ['平均单专区收益', '37,683.92元', '-'],
    ['平均单专区成本', '13,531.82元', '-'],
    ['成本超标专区', '10个', '总成本>32,000元'],
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

# 主要问题
doc.add_heading('【主要问题】', 2)
problems = [
    '河北、内蒙古、辽宁、吉林等省份成本率严重超标（>100%）',
    '10个专区总成本超过32,000元基线',
    '部分专区零收益，成本无法回收',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

# 建议措施
doc.add_heading('【建议措施】', 2)
measures = [
    '对成本率超标省份（河北、内蒙古、辽宁、吉林）进行成本复盘',
    '对零收益专区进行下线评估',
    '对高收益专区（湖北、新疆）加大上量支持',
]
for m in measures:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# 二、核心数据总览
doc.add_heading('二、核心数据总览', 1)

# 2.1 各省份专区分布
doc.add_heading('2.1 各省份专区分布', 2)
province_stats = df_c.groupby('所属省份').agg({
    '合同编号': 'count',
    '总收益情况': 'sum',
    '总成本': 'sum'
}).reset_index()
province_stats.columns = ['省份', '专区数量', '累计收益', '累计成本']
province_stats['成本率'] = province_stats.apply(lambda x: f"{(x['累计成本'] / x['累计收益'] * 100):.1f}%" if x['累计收益'] > 0 else 'N/A', axis=1)
province_stats = province_stats.sort_values('累计收益', ascending=False)

table2 = doc.add_table(rows=len(province_stats)+1, cols=5)
table2.style = 'Light Grid Accent 1'
headers2 = ['省份', '专区数量', '累计收益', '累计成本', '成本率']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, row in enumerate(province_stats.itertuples(), 1):
    table2.rows[idx].cells[0].text = str(row.省份)
    table2.rows[idx].cells[1].text = str(int(row.专区数量))
    table2.rows[idx].cells[2].text = f"{row.累计收益:,.0f}"
    table2.rows[idx].cells[3].text = f"{row.累计成本:,.0f}"
    table2.rows[idx].cells[4].text = str(row.成本率)

# 2.2 TOP10收益专区
doc.add_heading('2.2 TOP10收益专区', 2)
top10 = df_c.nlargest(10, '总收益情况')[[' 原专区名称', '所属省份', '总成本', '总收益情况', '成本率']]

table3 = doc.add_table(rows=11, cols=5)
table3.style = 'Light Grid Accent 1'
headers3 = ['排名', '专区名称', '省份', '累计收益', '成本率']
for i, header in enumerate(headers3):
    table3.rows[0].cells[i].text = header
    table3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(top10.iterrows(), 1):
    table3.rows[idx].cells[0].text = str(idx)
    table3.rows[idx].cells[1].text = str(row[' 原专区名称'])[:20]
    table3.rows[idx].cells[2].text = str(row['所属省份'])
    table3.rows[idx].cells[3].text = f"{row['总收益情况']:,.0f}"
    cost_rate = row['成本率']
    table3.rows[idx].cells[4].text = f"{cost_rate:.1f}%" if cost_rate < 999999 else 'N/A'

# 2.3 TOP10成本专区
doc.add_heading('2.3 TOP10成本专区（需关注）', 2)
top10_cost = df_c.nlargest(10, '总成本')[[' 原专区名称', '所属省份', '总成本', '总收益情况']]

table4 = doc.add_table(rows=11, cols=5)
table4.style = 'Light Grid Accent 1'
headers4 = ['排名', '专区名称', '省份', '总成本', '总收益情况']
for i, header in enumerate(headers4):
    table4.rows[0].cells[i].text = header
    table4.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(top10_cost.iterrows(), 1):
    table4.rows[idx].cells[0].text = str(idx)
    table4.rows[idx].cells[1].text = str(row[' 原专区名称'])[:20]
    table4.rows[idx].cells[2].text = str(row['所属省份'])
    table4.rows[idx].cells[3].text = f"{row['总成本']:,.0f}"
    table4.rows[idx].cells[4].text = f"{row['总收益情况']:,.0f}" if row['总收益情况'] > 0 else '0'

doc.add_page_break()

# 三、数据分析与洞察
doc.add_heading('三、数据分析与洞察', 1)

# 3.1 成本分析
doc.add_heading('3.1 成本分析', 2)
cost_analysis = [
    '上线前平均成本：13,531.82元（远低于基线32,000元）',
    '成本超标专区：10个（总成本>32,000元）',
    '成本率超标省份：河北(347%)、内蒙古(180%)、辽宁(132%)、吉林(1761%)',
    '整体成本率：35.9%（健康水平）',
]
for item in cost_analysis:
    doc.add_paragraph(item, style='List Bullet')

# 3.2 收益分析
doc.add_heading('3.2 收益分析', 2)
revenue_analysis = [
    '平均单专区收益：37,683.92元',
    '收益最高专区：昌吉城建市政招采平台（767,865元）',
    'TOP3收益省份：湖北(135万)、新疆(130万)、河南(98万)',
    '零收益专区：需进一步排查原因',
]
for item in revenue_analysis:
    doc.add_paragraph(item, style='List Bullet')

# 3.3 问题专区清单
doc.add_heading('3.3 高风险专区清单', 2)
high_risk = df_c[(df_c['总成本'] > 32000) | (df_c['总收益情况'] == 0)]
doc.add_paragraph(f'共识别高风险专区 {len(high_risk)} 个，具体如下：')

table5 = doc.add_table(rows=min(len(high_risk), 15)+1, cols=4)
table5.style = 'Light Grid Accent 1'
headers5 = ['专区名称', '省份', '总成本', '总收益']
for i, header in enumerate(headers5):
    table5.rows[0].cells[i].text = header
    table5.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(high_risk.head(15).iterrows(), 1):
    table5.rows[idx].cells[0].text = str(row[' 原专区名称'])[:20]
    table5.rows[idx].cells[1].text = str(row['所属省份'])
    table5.rows[idx].cells[2].text = f"{row['总成本']:,.0f}"
    table5.rows[idx].cells[3].text = f"{row['总收益情况']:,.0f}"

doc.add_page_break()

# 四、重点工作规划
doc.add_heading('四、重点工作规划', 1)

# 4.1 成本管控
doc.add_heading('4.1 成本管控问题专区分析', 2)
doc.add_paragraph('责任人：XXX    完成时间：2026年4月30日')
doc.add_paragraph('问题专区清单：焦作市全要素阳光交易平台、张家口市第一医院招标采购平台、沈铁专区等10个专区')
doc.add_paragraph('成本复盘计划：')
plans = [
    '分析成本构成：人力成本、系统对接成本、运维成本占比',
    '识别超支原因：需求变更、实施周期延长、技术难点',
    '制定控制措施：严格需求变更管理、优化实施流程、加强进度监控',
]
for p in plans:
    doc.add_paragraph(p, style='List Bullet')

# 4.2 上量支持
doc.add_heading('4.2 上量潜力专区上量工作', 2)
doc.add_paragraph('责任人：XXX    完成时间：2026年4月30日')
doc.add_paragraph('潜力专区清单：昌吉城建市政招采平台、湖北专区、荆门城控专区等TOP10收益专区')
doc.add_paragraph('上量措施：')
measures2 = [
    '增加推广资源投入：配置专项推广预算',
    '优化用户体验：提升系统稳定性、简化操作流程',
    '开展营销活动：联合客户举办培训会、推介会',
]
for m in measures2:
    doc.add_paragraph(m, style='List Bullet')

# 保存
doc.save('D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\e交易专区收益成本统计报告_C合同_110个专区_完整版.docx')
print('完整报告已生成！')
print('报告包含：')
print('  - 执行摘要（关键指标、主要问题、建议措施）')
print('  - 核心数据总览（省份分布、TOP10收益、TOP10成本）')
print('  - 数据分析与洞察（成本分析、收益分析、高风险专区）')
print('  - 重点工作规划（成本管控、上量支持）')
