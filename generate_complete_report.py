# -*- coding: utf-8 -*-
"""
🍊 小橘生成完整版报告
包含：
1. 26年度新开设专区分析
2. 原有专区（非26年度）分析
   - 原有专区总收益TOP10
   - 原有专区成本TOP10
   - 26年度产生收益的专区排名
"""

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print('小橘开始生成完整版报告...')

# 读取数据
df_all = pd.read_excel('D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\专区信息汇总表_中原华北.xlsx', sheet_name=None)

all_data = []
for sheet_name, df in df_all.items():
    df['所属省份'] = sheet_name.split('_')[0]
    all_data.append(df)

df_combined = pd.concat(all_data, ignore_index=True)

# 数据预处理
df_combined['合同编号'] = df_combined['合同编号'].str.strip()
df_combined['确认接入时间'] = pd.to_datetime(df_combined['确认接入时间'], errors='coerce')
df_combined['接入年份'] = df_combined['确认接入时间'].dt.year
df_combined['总收益情况'] = pd.to_numeric(df_combined['总收益情况'], errors='coerce').fillna(0)
df_combined['专区成本'] = pd.to_numeric(df_combined['专区成本'], errors='coerce').fillna(0)

# 筛选C开头或暂无
df_filtered = df_combined[
    (df_combined['合同编号'].str.startswith('C', na=False)) | 
    (df_combined['合同编号'] == '暂无')
].copy()

# 区分26年度新开设和原有专区
df_2026 = df_filtered[df_filtered['接入年份'] == 2026].copy()
df_old = df_filtered[df_filtered['接入年份'] != 2026].copy()

print(f'26年度新开设: {len(df_2026)}个')
print(f'原有专区: {len(df_old)}个')
print(f'总计: {len(df_filtered)}个')

# ========== 生成可视化图表 ==========
plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

fig, axes = plt.subplots(2, 2, figsize=(16, 12))
fig.suptitle('e交易专区数据分析（完整版 - 123个专区）', fontsize=18, fontweight='bold')

# 图1: 原有专区收益TOP10
ax1 = axes[0, 0]
top10_old_revenue = df_old.nlargest(10, '总收益情况')
y_pos = np.arange(len(top10_old_revenue))
bars1 = ax1.barh(y_pos, top10_old_revenue['总收益情况']/10000, color='#2ecc71', alpha=0.8)
ax1.set_yticks(y_pos)
ax1.set_yticklabels([name[:8] for name in top10_old_revenue[' 原专区名称']], fontsize=9)
ax1.set_xlabel('累计收益（万元）', fontsize=11)
ax1.set_title('原有专区收益TOP10', fontsize=13, fontweight='bold')
ax1.invert_yaxis()
ax1.grid(axis='x', alpha=0.3)

# 图2: 原有专区成本TOP10
ax2 = axes[0, 1]
top10_old_cost = df_old.nlargest(10, '专区成本')
colors = ['#e74c3c' if cost > 32000 else '#f39c12' for cost in top10_old_cost['专区成本']]
bars2 = ax2.barh(np.arange(len(top10_old_cost)), top10_old_cost['专区成本']/10000, color=colors, alpha=0.8)
ax2.set_yticks(np.arange(len(top10_old_cost)))
ax2.set_yticklabels([name[:8] for name in top10_old_cost[' 原专区名称']], fontsize=9)
ax2.set_xlabel('专区成本（万元）', fontsize=11)
ax2.set_title('原有专区成本TOP10（红色=超标）', fontsize=13, fontweight='bold')
ax2.axvline(x=3.2, color='red', linestyle='--', alpha=0.7, label='基线3.2万')
ax2.invert_yaxis()
ax2.legend()
ax2.grid(axis='x', alpha=0.3)

# 图3: 26年度专区成本分布
ax3 = axes[1, 0]
if len(df_2026) > 0:
    cost_ranges = ['≤5千', '5千-1万', '1万-2万', '2万-3.2万', '>3.2万']
    cost_counts = [
        len(df_2026[df_2026['专区成本'] <= 5000]),
        len(df_2026[(df_2026['专区成本'] > 5000) & (df_2026['专区成本'] <= 10000)]),
        len(df_2026[(df_2026['专区成本'] > 10000) & (df_2026['专区成本'] <= 20000)]),
        len(df_2026[(df_2026['专区成本'] > 20000) & (df_2026['专区成本'] <= 32000)]),
        len(df_2026[df_2026['专区成本'] > 32000])
    ]
    colors3 = ['#2ecc71', '#27ae60', '#f39c12', '#e67e22', '#e74c3c']
    bars3 = ax3.bar(cost_ranges, cost_counts, color=colors3, alpha=0.8)
    ax3.set_xlabel('成本区间', fontsize=11)
    ax3.set_ylabel('专区数量', fontsize=11)
    ax3.set_title('26年度新开设专区成本分布', fontsize=13, fontweight='bold')
    ax3.grid(axis='y', alpha=0.3)
    for bar in bars3:
        height = bar.get_height()
        if height > 0:
            ax3.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}', 
                     ha='center', va='bottom', fontsize=10, fontweight='bold')

# 图4: 26年度收益分析（虽然都是0，但展示结构）
ax4 = axes[1, 1]
ax4.text(0.5, 0.5, '26年度新开设专区\n目前收益均为0\n需关注后续上量', 
         ha='center', va='center', fontsize=14, 
         bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
ax4.set_xlim(0, 1)
ax4.set_ylim(0, 1)
ax4.axis('off')
ax4.set_title('26年度收益情况', fontsize=13, fontweight='bold')

plt.tight_layout()
plt.savefig('D:\\openclaw-workspace\\e交易专区分析图表_完整版.png', dpi=150, bbox_inches='tight')
print('图表生成完成')

# ========== 生成Word报告 ==========
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

# 封面
title = doc.add_heading('e交易专区收益成本统计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('（合同编号C开头或暂无 - 123个专区完整分析）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
doc.add_paragraph('中原区、华北区运营分析报告').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('报告日期：2026年3月31日').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)
doc.add_paragraph(f'截至2026年3月31日，新点e交易（中原、华北区）合同编号以C开头或暂无的专区共运营{len(df_filtered)}个。')
doc.add_paragraph(f'其中：26年度新开设{len(df_2026)}个，原有专区{len(df_old)}个。')
doc.add_paragraph('【重要说明】专区成本已包含上线前成本，不再重复计算。')

# 关键指标
doc.add_heading('【关键指标】', 2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['指标名称', '数值', '说明']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

total_revenue = df_filtered['总收益情况'].sum()
total_cost = df_filtered['专区成本'].sum()
avg_cost = df_filtered['专区成本'].mean()
high_cost_count = len(df_filtered[df_filtered['专区成本'] > 32000])

data = [
    ['专区总数', f'{len(df_filtered)}个', 'C开头或暂无'],
    ['原有专区', f'{len(df_old)}个', '2025年及以前接入'],
    ['26年度新开设', f'{len(df_2026)}个', '2026年接入'],
    ['累计收益', f'{total_revenue:,.2f}元', '-'],
    ['累计成本', f'{total_cost:,.2f}元', '专区成本'],
    ['平均成本', f'{avg_cost:,.2f}元', '专区成本'],
    ['成本超标专区', f'{high_cost_count}个', '>32,000元'],
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_page_break()

# 二、原有专区分析（2025年及以前）
doc.add_heading('二、原有专区分析（2025年及以前接入）', 1)
doc.add_paragraph(f'原有专区共{len(df_old)}个，是收益贡献的主力。')

# 2.1 原有专区收益TOP10
doc.add_heading('2.1 原有专区总收益TOP10', 2)
top10_revenue = df_old.nlargest(10, '总收益情况')[[' 原专区名称', '所属省份', '专区成本', '总收益情况']]

table2 = doc.add_table(rows=11, cols=5)
table2.style = 'Light Grid Accent 1'
headers2 = ['排名', '专区名称', '省份', '累计收益', '专区成本']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(top10_revenue.iterrows(), 1):
    table2.rows[idx].cells[0].text = str(idx)
    table2.rows[idx].cells[1].text = str(row[' 原专区名称'])[:18]
    table2.rows[idx].cells[2].text = str(row['所属省份'])
    table2.rows[idx].cells[3].text = f'{row["总收益情况"]:,.0f}'
    table2.rows[idx].cells[4].text = f'{row["专区成本"]:,.0f}'

# 2.2 原有专区成本TOP10
doc.add_heading('2.2 原有专区成本TOP10（需重点关注）', 2)
top10_cost = df_old.nlargest(10, '专区成本')[[' 原专区名称', '所属省份', '专区成本', '总收益情况']]

table3 = doc.add_table(rows=11, cols=6)
table3.style = 'Light Grid Accent 1'
headers3 = ['排名', '专区名称', '省份', '专区成本', '累计收益', '状态']
for i, header in enumerate(headers3):
    table3.rows[0].cells[i].text = header
    table3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(top10_cost.iterrows(), 1):
    table3.rows[idx].cells[0].text = str(idx)
    table3.rows[idx].cells[1].text = str(row[' 原专区名称'])[:18]
    table3.rows[idx].cells[2].text = str(row['所属省份'])
    table3.rows[idx].cells[3].text = f'{row["专区成本"]:,.0f}'
    table3.rows[idx].cells[4].text = f'{row["总收益情况"]:,.0f}'
    if row['专区成本'] > 32000:
        table3.rows[idx].cells[5].text = '超标'
    else:
        table3.rows[idx].cells[5].text = '正常'