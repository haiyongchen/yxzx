#!/usr/bin/env python3
"""
e交易专区收益成本统计报告生成器（增强版）
支持图表生成和新增指标
"""

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os

plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def analyze_data(file_path):
    df = pd.read_excel(file_path)
    df['确认接入时间'] = pd.to_datetime(df['确认接入时间'], errors='coerce')
    
    current_date = pd.to_datetime('2026-04-08')
    half_year_ago = current_date - timedelta(days=180)
    
    stats = {
        'total': len(df),
        'indicator1': len(df[(df['确认接入时间'] < '2025-04-08') & (df['总收益情况'] == 0)]),
        'indicator2': len(df[(df['确认接入时间'] < '2025-04-08') & (df['总收益情况'] > 0) & (df['总收益情况'] < 100000)]),
        'indicator3': len(df[(df['确认接入时间'] < '2025-04-08') & (df['总收益情况'] > 0) & (df['总收益情况'] < 50000)]),
        'indicator4': len(df[(df['确认接入时间'] < '2025-01-01') & (df['25年收益情况'] > 0) & (df['25年收益情况'] < 100000)]),
        'indicator5': len(df[(df['确认接入时间'] < '2025-01-01') & (df['25年收益情况'] > 0) & (df['25年收益情况'] < 50000)]),
        'indicator6': len(df[df['26年收益情况'] > 0]),
        'indicator7': len(df[(df['25年收益情况'] > 0) & (df['26年收益情况'] == 0)]),
        'indicator8': len(df[(df['确认接入时间'] < half_year_ago) & (df['总收益情况'] == 0)]),
    }
    
    risk_data = {
        '红色-重点关注': stats['indicator1'] + stats['indicator2'],
        '橙色-需改进': stats['indicator4'],
        '黄色-观察': stats['indicator6'],
        '灰色-流失风险': stats['indicator7'],
    }
    
    df_25 = df[(df['确认接入时间'] < '2025-01-01') & (df['25年收益情况'] > 0) & (df['25年收益情况'] < 100000)]
    revenue_bins = [0, 30000, 50000, 70000, 100000]
    revenue_labels = ['0-3万', '3-5万', '5-7万', '7-10万']
    df_25['收益区间'] = pd.cut(df_25['25年收益情况'], bins=revenue_bins, labels=revenue_labels)
    revenue_dist = df_25['收益区间'].value_counts().sort_index()
    
    df_loss = df[(df['25年收益情况'] > 0) & (df['26年收益情况'] == 0)].copy()
    df_loss = df_loss.sort_values('25年收益情况', ascending=False).head(10)
    
    return stats, risk_data, revenue_dist, df_loss, df

def create_chart1(risk_data, output_path):
    colors = ['#FF6B6B', '#FFA500', '#FFD700', '#808080']
    fig, ax = plt.subplots(figsize=(10, 6))
    wedges, texts, autotexts = ax.pie(
        risk_data.values(), 
        labels=risk_data.keys(),
        autopct='%1.1f%%',
        colors=colors,
        startangle=90,
        textprops={'fontsize': 12}
    )
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')
    ax.set_title('e交易专区风险等级分布', fontsize=16, fontweight='bold', pad=20)
    ax.legend(wedges, [f'{k}: {v}个' for k, v in risk_data.items()],
              title="风险等级", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'图表1已生成: {output_path}')

def create_chart2(revenue_dist, output_path):
    fig, ax = plt.subplots(figsize=(10, 6))
    bars = ax.bar(revenue_dist.index, revenue_dist.values, color='#4A90D9', edgecolor='black')
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}',
                ha='center', va='bottom', fontsize=11)
    ax.set_xlabel('25年收益区间', fontsize=12)
    ax.set_ylabel('专区数量', fontsize=12)
    ax.set_title('25年前接入专区收益分布（25年总收益<10万）', fontsize=14, fontweight='bold')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'图表2已生成: {output_path}')

def create_chart3(df_loss, output_path):
    fig, ax = plt.subplots(figsize=(12, 6))
    y_pos = range(len(df_loss))
    bars = ax.barh(y_pos, df_loss['25年收益情况'].values, color='#E74C3C', edgecolor='black')
    ax.set_yticks(y_pos)
    ax.set_yticklabels(df_loss['专区名称'].values, fontsize=10)
    for i, bar in enumerate(bars):
        width = bar.get_width()
        ax.text(width, bar.get_y() + bar.get_height()/2., f' {int(width):,}',
                ha='left', va='center', fontsize=9)
    ax.set_xlabel('25年收益（元）', fontsize=12)
    ax.set_title('TOP10 流失风险专区（25年有收益但26年无）', fontsize=14, fontweight='bold')
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    ax.invert_yaxis()
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'图表3已生成: {output_path}')

def create_word_report(stats, output_path, chart_paths):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    title = doc.add_heading('新点e交易专区（中原/华北）收益成本', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('近期基于新点e交易平台各专区的运营数据，从收益、成本、时间维度进行深度分析。通过部分指标的统计分析，我们发现当前平台存在一定数量的零收益、低收益专区，需要引起重点关注并采取针对性措施。')
    doc.add_paragraph(f'数据统计范围涵盖中原、华北等区域共{stats["total"]}个专区，重点分析确认接入时间、25年总收益、26年总收益、总成本等关键指标。通过建立风险分级体系（红色-重点关注、橙色-需改进、黄色-观察、灰色-流失风险）。')
    doc.add_paragraph(f'本次统计识别出一批需重点关注的低收益专区，其中纳入红色等级预警的有 {stats["indicator1"] + stats["indicator2"]} 个、橙色等级预警的有 {stats["indicator4"]} 个（两类等级存在部分专区重叠），另有灰色等级 {stats["indicator7"]} 个。从数据趋势来看，部分早期接入的专区已出现运营效率低下、客户活跃度下降等问题。')
    doc.add_paragraph('对于长期无收益且无可挽救价值的专区，建议果断下线以节约运营成本；对于有潜力的专区，加大资源投入和运营支持。')
    
    doc.add_heading('一、核心数据概览', level=1)
    doc.add_heading('1.1 七大指标统计', level=2)
    
    doc.add_paragraph(f'从上表可以看出，接入超1年，但长期零收益专区为0的有{stats["indicator1"]}个，长期低收益专区（总收益10w以下）共计{stats["indicator2"]}个，占比较高，因收益10w以下的专区过多，故进一步分析收益5w以下的专区，数量为{stats["indicator3"]}个；25年前接入但收益不达标的专区（25年总收益在10w以下）共计{stats["indicator4"]}个，其中25年总收益在5w以下的有{stats["indicator5"]}个，反映出早期接入专区的运营效率问题；特别值得关注的是，有{stats["indicator7"]}个专区在25年产生收益但26年暂无收益，存在客户流失风险；此外，接入超过半年但仍无收益的专区有{stats["indicator8"]}个，需要重点关注。')
    
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    headers = ['指标', '条件', '数量', '风险等级', '备注']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data_rows = [
        ['指标一', '接入超过1年，总收益为0', str(stats['indicator1']), '红色', '重点关注'],
        ['指标二', '接入超过1年，0<总收益<10w', str(stats['indicator2']), '红色', '重点关注'],
        ['指标三', '接入超过1年，0<总收益<5w', str(stats['indicator3']), '红色', '重点关注'],
        ['指标四', '25年前接入，0<25年收益<10w', str(stats['indicator4']), '橙色', '需改进'],
        ['指标五', '25年前接入，0<25年收益<5w', str(stats['indicator5']), '橙色', '需改进'],
        ['指标六', '26年产生收益', str(stats['indicator6']), '黄色', '观察'],
        ['指标七', '25年有收益，26年无', str(stats['indicator7']), '灰色', '流失风险'],
        ['指标八', '接入超半年，总收益为0', str(stats['indicator8']), '深红', '紧急关注'],
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_heading('1.2 风险等级分布', level=2)
    doc.add_paragraph('根据收益水平、接入时长、趋势变化等维度，我们将专区划分为四个风险等级：')
    
    doc.add_paragraph('【红色-重点关注】', style='Heading 3')
    doc.add_paragraph('处理对象：接入超过一年，但长期零收益或者低收益专区')
    doc.add_paragraph('处理措施：')
    doc.add_paragraph('逐一排查专区运营状态，确认是否仍在正常服务客户', style='List Bullet')
    doc.add_paragraph('对于长期无交易的专区，评估是否继续投入运营成本', style='List Bullet')
    doc.add_paragraph('制定专区下线计划，释放运维资源', style='List Bullet')
    doc.add_paragraph('对于仍有潜力的专区，制定专项提升计划，明确责任人和时间节点', style='List Bullet')
    
    doc.add_paragraph('【橙色-需改进】', style='Heading 3')
    doc.add_paragraph('处理对象：25年前接入，25年存在收益但小于10w')
    doc.add_paragraph('处理措施：')
    doc.add_paragraph('分析专区情况，找出收益低的原因，是否客户业务量上限，是否存在上量可能', style='List Bullet')
    
    doc.add_paragraph('【黄色-观察】', style='Heading 3')
    doc.add_paragraph('26年产生收益的专区，运营状态良好，继续保持观察。')
    
    doc.add_paragraph('【灰色-流失风险】', style='Heading 3')
    doc.add_paragraph('处理对象：25年有收益但26年无收益专区')
    doc.add_paragraph('处理措施：')
    doc.add_paragraph('优先安排人员对该部分专区进行情况了解', style='List Bullet')
    doc.add_paragraph('了解客户流失原因', style='List Bullet')
    
    # 插入图1
    doc.add_paragraph()
    doc.add_picture(chart_paths[0], width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 二、收益分析
    doc.add_heading('二、收益分析', level=1)
    doc.add_heading('2.1 25年收益分布', level=2)
    doc.add_paragraph('针对25年前接入的低收益专区，我们进一步分析其25年总收益分布情况。从数据来看，收益分布呈现明显的两极分化特征：大部分专区集中在0-3万元区间，而收益在7-10万元区间的专区数量相对较少。这一分布特征表明，当前低收益专区普遍存在运营效率不高的问题，需要系统性分析原因并制定提升策略。')
    
    # 插入图2
    doc.add_paragraph()
    doc.add_picture(chart_paths[1], width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2.2 流失风险专区TOP10', level=2)
    doc.add_paragraph('以下10个专区在2025年产生较高收益，但2026年至今暂无收益记录，存在较高的客户流失或业务停滞风险，建议优先安排客户经理进行回访和原因排查。')
    
    # 插入图3
    doc.add_paragraph()
    doc.add_picture(chart_paths[2], width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 三、后续处理建议
    doc.add_heading('三、后续处理建议', level=1)
    doc.add_paragraph('基于以上数据分析，针对不同类型的低收益专区，我们提出以下分类处理建议：')
    
    doc.add_heading('3.1 红色等级专区（立即处理）', level=2)
    doc.add_paragraph('处理对象：接入超过一年，但长期零收益或者低收益专区')
    doc.add_paragraph('处理措施：')
    doc.add_paragraph('逐一排查专区运营状态，确认是否仍在正常服务客户', style='List Bullet')
    doc.add_paragraph('对于长期无交易的专区，评估是否继续投入运营成本', style='List Bullet')
    doc.add_paragraph('制定专区下线计划，释放运维资源', style='List Bullet')
    doc.add_paragraph('对于仍有潜力的专区，制定专项提升计划，明确责任人和时间节点', style='List Bullet')
    
    doc.add_heading('3.2 灰色等级专区（紧急跟进）', level=2)
    doc.add_paragraph('处理对象：25年有收益但26年无收益专区')
    doc.add_paragraph('处理措施：')
    doc.add_paragraph('优先安排人员对该部分专区进行情况了解', style='List Bullet')
    doc.add_paragraph('了解客户流失原因', style='List Bullet')
    
    doc.add_heading('3.3 橙色等级专区（持续优化）', level=2)
    doc.add_paragraph('处理对象：25年前接入，25年存在收益但小于10w')
    doc.add_paragraph('处理措施：')
    doc.add_paragraph('分析专区情况，找出收益低的原因，是否客户业务量上限，是否存在上量可能', style='List Bullet')
    
    # 四、关键任务
    doc.add_heading('四、关键任务', level=1)
    
    table2 = doc.add_table(rows=5, cols=5)
    table2.style = 'Table Grid'
    headers2 = ['序号', '实施工作', '负责人', '计划完成节点', '备注']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    task_rows = [
        ['1', '对25年有收益但26年未产生收益的大企业侧专区进行摸排', '陈海勇/钟明珠', '2026年4月17日', ''],
        ['2', '对于接入超一年的低收益的大企业侧专区逐一排查，看是否存在上量可能', '陈海勇/钟明珠', '2026年4月30日', ''],
        ['3', '对接入超一年的零收益专区摸排，看是否存在上量可能，确认是否需要执行下线操作', '陈海勇/钟明珠', '2026年5月31日', ''],
        ['4', '对于26年产生收益专区持续关注', '陈海勇/钟明珠', '2026年5月31日', ''],
    ]
    
    for i, row_data in enumerate(task_rows, 1):
        for j, cell_data in enumerate(row_data):
            table2.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('【附录】').runs[0].font.bold = True
    doc.add_paragraph('详细数据请参见附件《中原华北区专区数据_成本更新.xlsx》')
    
    doc.save(output_path)
    print(f'Word报告已生成: {output_path}')

def main():
    data_file = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\e交易数据分析处理文件\\中原华北区专区数据_成本更新.xlsx'
    output_dir = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\e交易数据分析处理文件'
    
    chart1_path = os.path.join(output_dir, 'chart1_risk_distribution.png')
    chart2_path = os.path.join(output_dir, 'chart2_revenue_distribution.png')
    chart3_path = os.path.join(output_dir, 'chart3_top10_loss.png')
    report_path = os.path.join(output_dir, 'e交易专区收益成本统计报告_完整版.docx')
    
    print('=== 开始生成报告 ===')
    print('1. 分析数据...')
    stats, risk_data, revenue_dist, df_loss, df = analyze_data(data_file)
    
    print(f'   总专区数: {stats["total"]}')
    print(f'   指标一: {stats["indicator1"]}')
    print(f'   指标二: {stats["indicator2"]}')
    print(f'   指标三: {stats["indicator3"]}')
    print(f'   指标四: {stats["indicator4"]}')
    print(f'   指标五: {stats["indicator5"]}')
    print(f'   指标六: {stats["indicator6"]}')
    print(f'   指标七: {stats["indicator7"]}')
    print(f'   指标八(新增): {stats["indicator8"]}')
    
    print('2. 生成图表...')
    create_chart1(risk_data, chart1_path)
    create_chart2(revenue_dist, chart2_path)
    create_chart3(df_loss, chart3_path)
    
    print('3. 生成Word报告...')
    create_word_report(stats, report_path, [chart1_path, chart2_path, chart3_path])
    
    print('=== 报告生成完成 ===')

if __name__ == '__main__':
    main()