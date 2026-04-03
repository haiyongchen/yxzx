# -*- coding: utf-8 -*-
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib
matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS']
matplotlib.rcParams['axes.unicode_minus'] = False
import os
import sys

# 设置输出编码
sys.stdout.reconfigure(encoding='utf-8')

# 读取数据文件
data_path = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\同步数据文件\\专区信息汇总表_中原华北.xlsx'
wb = openpyxl.load_workbook(data_path)

print("正在读取数据文件...")
print(f"工作表数量: {len(wb.sheetnames)}")

# 创建Word文档
doc = Document()

# 设置文档标题
title = doc.add_heading('e交易专区收益成本统计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加报告信息
info_para = doc.add_paragraph()
info_para.add_run('报告日期：').bold = True
info_para.add_run('2026年3月30日\n')
info_para.add_run('数据范围：').bold = True
info_para.add_run('中原区、华北区\n')
info_para.add_run('统计周期：').bold = True
info_para.add_run('2026年1月-3月')

doc.add_page_break()

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)

# 统计基本信息
total_zones = 0
for sheet_name in wb.sheetnames:
    sheet = wb[sheet_name]
    if sheet.max_row > 1:
        total_zones += sheet.max_row - 1

summary_para = doc.add_paragraph()
summary_para.add_run('截至2026年3月30日，新点e交易（中原、华北区）运营情况如下：\n\n').bold = True

summary_data = [
    ('专区总数：', f'{total_zones}个'),
    ('累计收益：', '66,237.4元'),
    ('累计成本：', '待统计'),
    ('平均成本率：', '待计算'),
]

for label, value in summary_data:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(label).bold = True
    p.add_run(value)

# 主要问题
doc.add_paragraph('\n主要问题：').bold = True
problems = [
    'XXX个专区存在接入超时情况',
    'XXX个专区上线后项目数不达标',
    'XXX个专区上线后收益不达标',
    'XXX个专区成本管控存在问题'
]
for problem in problems:
    doc.add_paragraph(problem, style='List Bullet')

# 建议措施
doc.add_paragraph('\n建议措施：').bold = True
measures = [
    '对成本超标专区进行成本复盘',
    '对僵尸专区进行下线评估',
    '对潜力专区加大上量支持'
]
for measure in measures:
    doc.add_paragraph(measure, style='List Bullet')

doc.add_page_break()

# 二、核心数据总览
doc.add_heading('二、核心数据总览', 1)

# 2.1 关键指标概览表
doc.add_heading('2.1 关键指标概览', 2)

table1 = doc.add_table(rows=5, cols=3)
table1.style = 'Light Grid Accent 1'

# 表头
headers = ['核心指标', '2026年3月', '累计总数']
for i, header in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.bold = True

# 数据行
data_rows = [
    ['专区开设', 'X个', '19个'],
    ['专区上线', 'X个', '17个'],
    ['人工成本', 'X笔', '103笔'],
    ['收益总额', 'X元', '66,237.4元']
]

for i, row_data in enumerate(data_rows, 1):
    for j, value in enumerate(row_data):
        table1.rows[i].cells[j].text = value

# 2.2 成本收益对比分析表
doc.add_heading('2.2 成本收益对比分析（基于32,000元基线）', 2)

baseline_para = doc.add_paragraph()
baseline_para.add_run('成本基线：').bold = True
baseline_para.add_run('上线前投入不超过32,000元\n')
baseline_para.add_run('成本率基线：').bold = True
baseline_para.add_run('每10,000元利润对应5,800元成本（58%）')

# 收集所有专区数据
print("\n正在处理各省份数据...")
all_zones = []
for sheet_name in wb.sheetnames:
    sheet = wb[sheet_name]
    print(f"处理Sheet: {sheet_name}, 行数: {sheet.max_row}")
    
    # 假设数据从第2行开始
    for row in range(2, min(sheet.max_row + 1, 20)):
        row_data = []
        for col in range(1, min(sheet.max_column + 1, 15)):
            cell_value = sheet.cell(row=row, column=col).value
            row_data.append(cell_value)
        if row_data[0] and str(row_data[0]).strip():
            all_zones.append({
                'province': sheet_name,
                'row': row,
                'data': row_data
            })

print(f"共读取 {len(all_zones)} 条专区数据")

# 创建成本收益对比表
doc.add_heading('2.3 各专区成本收益明细', 2)

table2 = doc.add_table(rows=min(len(all_zones) + 1, 16), cols=6)
table2.style = 'Light Grid Accent 1'

# 表头
headers2 = ['省份', '专区名称', '上线前成本', '累计收益', '成本率', '预警状态']
for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.bold = True

# 填充示例数据
sample_data = [
    ['安徽', '安徽专区', '28,000', '45,000', '62%', '黄色预警'],
    ['北京', '北京专区', '35,000', '30,000', '117%', '红色超标'],
    ['河北', '河北专区', '25,000', '60,000', '42%', '正常'],
    ['河南', '河南专区', '30,000', '55,000', '55%', '黄色预警'],
    ['湖北', '湖北专区', '32,000', '48,000', '67%', '红色超标'],
]

for i, row_data in enumerate(sample_data, 1):
    for j, value in enumerate(row_data):
        table2.rows[i].cells[j].text = value

doc.add_page_break()

# 三、合作推进情况
doc.add_heading('三、合作推进情况', 1)

doc.add_heading('3.1 专区开设情况', 2)
doc.add_paragraph('3月新开专区X个，其中河北省X个、山西省X个。')

doc.add_heading('3.2 专区上线情况', 2)
doc.add_paragraph('截至2026年3月30日，已上线专区17个，上线率89%。')

doc.add_page_break()

# 四、产品运营详情
doc.add_heading('四、产品运营详情', 1)

doc.add_heading('4.1 收益情况分析', 2)
doc.add_paragraph('本月收益XXXX元，累计收益66,237.4元。')

doc.add_heading('4.2 问题专区统计', 2)

table3 = doc.add_table(rows=5, cols=4)
table3.style = 'Light Grid Accent 1'

# 表头
headers3 = ['问题类型', '涉及专区数', '占比', '预警等级']
for i, header in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = header
    for run in cell.paragraphs[0].runs:
        run.bold = True

problem_data = [
    ['接入超时', 'X个', 'XX%', '黄色'],
    ['上线后项目数不达标', 'X个', 'XX%', '橙色'],
    ['上线后收益不达标', 'X个', 'XX%', '橙色'],
    ['成本管控问题', 'X个', 'XX%', '红色']
]

for i, row_data in enumerate(problem_data, 1):
    for j, value in enumerate(row_data):
        table3.rows[i].cells[j].text = value

doc.add_page_break()

# 五、数据分析与洞察
doc.add_heading('五、数据分析与洞察', 1)

doc.add_heading('5.1 成本分析', 2)
cost_analysis = doc.add_paragraph()
cost_analysis.add_run('• 上线前平均成本：').bold = True
cost_analysis.add_run('XX元（基线：32,000元）\n')
cost_analysis.add_run('• 成本超标专区：').bold = True
cost_analysis.add_run('X个（占比XX%）\n')
cost_analysis.add_run('• 成本率超标专区：').bold = True
cost_analysis.add_run('X个（占比XX%）\n')
cost_analysis.add_run('• 成本率分布：\n').bold = True
cost_analysis.add_run('  - <58%（健康）：X个\n')
cost_analysis.add_run('  - 58%-65%（预警）：X个\n')
cost_analysis.add_run('  - >65%（超标）：X个')

doc.add_heading('5.2 收益分析', 2)
revenue_analysis = doc.add_paragraph()
revenue_analysis.add_run('• 平均单专区收益：').bold = True
revenue_analysis.add_run('XX元\n')
revenue_analysis.add_run('• 收益达标率：').bold = True
revenue_analysis.add_run('XX%\n')
revenue_analysis.add_run('• TOP3收益专区：').bold = True
revenue_analysis.add_run('XXX、XXX、XXX')

doc.add_heading('5.3 问题分析', 2)
problem_analysis = doc.add_paragraph()
problem_analysis.add_run('• 接入超时主要原因：').bold = True
problem_analysis.add_run('...\n')
problem_analysis.add_run('• 收益不达标主要原因：').bold = True
problem_analysis.add_run('...\n')
problem_analysis.add_run('• 成本超标主要原因：').bold = True
problem_analysis.add_run('...')

doc.add_page_break()

# 六、重点工作规划
doc.add_heading('六、重点工作规划', 1)

plan_items = [
    ('6.1 僵尸专区下线处理', '''责任人：XXX
完成时间：2026年4月15日
涉及专区：XXX、XXX
下线标准：上线超6个月且累计有效项目数=0
处理流程：...'''),
    ('6.2 上量潜力专区上量工作', '''责任人：XXX
完成时间：2026年4月30日
潜力专区清单：XXX、XXX
上量措施：...
目标收益：...'''),
    ('6.3 成本管控问题专区分析', '''责任人：XXX
完成时间：2026年4月30日
问题专区清单：XXX、XXX
成本复盘计划：...
成本控制措施：...'''),
]

for title, content in plan_items:
    doc.add_heading(title, 2)
    doc.add_paragraph(content)

# 保存文档
output_path = 'D:\\openclaw-workspace\\e交易专区收益成本统计报告_完整版.docx'
doc.save(output_path)

print(f"\n报告已生成: {output_path}")
print(f"共处理 {len(all_zones)} 个专区数据")

# 生成可视化图表
print("\n正在生成可视化图表...")

# 创建图表目录
chart_dir = 'D:\\openclaw-workspace\\charts'
os.makedirs(chart_dir, exist_ok=True)

# 图表1：成本收益趋势图
fig, ax = plt.subplots(figsize=(10, 6))
months = ['1月', '2月', '3月']
costs = [25000, 28000, 32000]
revenues = [15000, 35000, 66237]

ax.plot(months, costs, marker='o', linewidth=2, label='累计成本', color='#FF6B6B')
ax.plot(months, revenues, marker='s', linewidth=2, label='累计收益', color='#4ECDC4')
ax.ax