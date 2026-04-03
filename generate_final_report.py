# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print('小橘开始生成最终版报告...')

# 读取数据
df = pd.read_csv('D:\\openclaw-workspace\\all_zones_c_contract.csv', encoding='utf-8-sig')

# 数据处理
df['确认接入时间'] = pd.to_datetime(df['确认接入时间'], errors='coerce')
df['接入年份'] = df['确认接入时间'].dt.year
df['总收益情况'] = pd.to_numeric(df['总收益情况'], errors='coerce').fillna(0)
df['专区成本'] = pd.to_numeric(df['专区成本'], errors='coerce').fillna(0)

# 区分26年度和原有专区
df_2026 = df[df['接入年份'] == 2026].copy()
df_old = df[df['接入年份'] != 2026].copy()

total_revenue = df['总收益情况'].sum()
total_cost = df['专区成本'].sum()
avg_cost = df['专区成本'].mean()
high_cost = len(df[df['专区成本'] > 32000])

# 创建文档
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

# 标题
title = doc.add_heading('e交易专区收益成本统计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('（合同编号C开头或暂无 - 专区名称维度分析）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
doc.add_paragraph('中原区、华北区运营分析报告').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('报告日期：2026年3月31日').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)
doc.add_paragraph('截至2026年3月31日，新点e交易（中原、华北区）合同编号以C开头或暂无的专区共运营' + str(len(df)) + '个，其中原有专区' + str(len(df_old)) + '个，26年度新开设' + str(len(df_2026)) + '个。')

# 关键指标
doc.add_heading('【关键指标】', 2)
table1 = doc.add_table(rows=7, cols=3)
table1.style = 'Light Grid Accent 1'
headers = ['指标名称', '数值', '说明']
for i, header in enumerate(headers):
    table1.rows[0].cells[i].text = header
    table1.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

table1.rows[1].cells[0].text = '专区总数'
table1.rows[1].cells[1].text = str(len(df)) + '个'
table1.rows[1].cells[2].text = 'C开头或暂无'

table1.rows[2].cells[0].text = '原有专区'
table1.rows[2].cells[1].text = str(len(df_old)) + '个'
table1.rows[2].cells[2].text = '2025年及以前接入'

table1.rows[3].cells[0].text = '26年度新开设'
table1.rows[3].cells[1].text = str(len(df_2026)) + '个'
table1.rows[3].cells[2].text = '2026年接入'

table1.rows[4].cells[0].text = '累计收益'
table1.rows[4].cells[1].text = str(int(total_revenue)) + '元'
table1.rows[4].cells[2].text = '-'

table1.rows[5].cells[0].text = '平均专区成本'
table1.rows[5].cells[1].text = str(int(avg_cost)) + '元'
table1.rows[5].cells[2].text = '专区成本'

table1.rows[6].cells[0].text = '成本超标专区'
table1.rows[6].cells[1].text = str(high_cost) + '个'
table1.rows[6].cells[2].text = '>32,000元'

# 主要问题
doc.add_heading('【主要问题】', 2)
doc.add_paragraph(str(high_cost) + '个专区成本超过32,000元基线', style='List Bullet')
doc.add_paragraph('部分专区零收益，成本无法回收', style='List Bullet')
doc.add_paragraph('26年度新开设' + str(len(df_2026)) + '个专区目前收益均为0，需关注后续上量', style='List Bullet')

# 建议措施
doc.add_heading('【建议措施】', 2)
doc.add_paragraph('对成本超标专区进行成本复盘', style='List Bullet')
doc.add_paragraph('对零收益专区进行下线评估', style='List Bullet')
doc.add_paragraph('对26年度新开设专区进行月度跟踪', style='List Bullet')

doc.add_page_break()

# 二、核心数据总览
doc.add_heading('二、核心数据总览', 1)

# 2.1 原有专区收益TOP10
doc.add_heading('2.1 原有专区收益TOP10（按专区名称维度）', 2)
doc.add_paragraph('【数据说明】专区名称取自各表格C列（原专区名称字段）')

top10_revenue = df_old.nlargest(10, '总收益情况')
table2 = doc.add_table(rows=11, cols=6)
table2.style = 'Light Grid Accent 1'
headers2 = ['专区名称', '专区成本', '累计收益', '成本率', '预警状态', '建议措施']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(top10_revenue.iterrows(), 1):
    table2.rows[idx].cells[0].text = str(row[' 原专区名称'])[:18]
    table2.rows[idx].cells[1].text = str(int(row['专区成本']))
    table2.rows[idx].cells[2].text = str(int(row['总收益情况']))
    
    if row['总收益情况'] > 0:
        cost_rate = row['专区成本'] / row['总收益情况'] * 100
        table2.rows[idx].cells[3].text = str(int(cost_rate)) + '%'
        
        if cost_rate > 65:
            table2.rows[idx].cells[4].text = '红色超标'
            table2.rows[idx].cells[5].text = '成本复盘约谈'
        elif cost_rate > 58:
            table2.rows[idx].cells[4].text = '黄色预警'
            table2.rows[idx].cells[5].text = '持续监控'
        else:
            table2.rows[idx].cells[4].text = '正常'
            table2.rows[idx].cells[5].text = '持续监控'
    else:
        table2.rows[idx].cells[3].text = 'N/A'
        table2.rows[idx].cells[4].text = '零收益'
        table2.rows[idx].cells[5].text = '下线评估'

# 2.2 原有专区成本TOP10
doc.add_heading('2.2 原有专区成本TOP10（需重点关注）', 2)
top10_cost = df_old.nlargest(10, '专区成本')

table3 = doc.add_table(rows=11, cols=6)
table3.style = 'Light Grid Accent 1'
for i, header in enumerate(headers2):
    table3.rows[0].cells[i].text = header
    table3.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(top10_cost.iterrows(), 1):
    table3.rows[idx].cells[0].text = str(row[' 原专区名称'])[:18]
    table3.rows[idx].cells[1].text = str(int(row['专区成本']))
    table3.rows[idx].cells[2].text = str(int(row['总收益情况']))
    
    if row['总收益情况'] > 0:
        cost_rate = row['专区成本'] / row['总收益情况'] * 100
        table3.rows[idx].cells[3].text = str(int(cost_rate)) + '%'
    else:
        table3.rows[idx].cells[3].text = 'N/A'
    
    if row['专区成本'] > 32000:
        table3.rows[idx].cells[4].text = '超标'
        table3.rows[idx].cells[5].text = '成本复盘约谈'
    else:
        table3.rows[idx].cells[4].text = '正常'
        table3.rows[idx].cells[5].text = '持续监控'

doc.add_page_break()

# 三、26年度新开设专区分析
doc.add_heading('三、26年度新开设专区分析', 1)

# 3.1 总体情况
doc.add_heading('3.1 总体情况', 2)
table_2026 = doc.add_table(rows=6, cols=3)
table_2026.style = 'Light Grid Accent 1'
headers_2026 = ['指标名称', '数值', '说明']
for i, header in enumerate(headers_2026):
    table_2026.rows[0].cells[i].text = header
    table_2026.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

table_2026.rows[1].cells[0].text = '新开设专区数'
table_2026.rows[1].cells[1].text = str(len(df_2026)) + '个'
table_2026.rows[1].cells[2].text = '2026年接入'

table_2026.rows[2].cells[0].text = '总成本'
table_2026.rows[2].cells[1].text = str(int(df_2026['专区成本'].sum())) + '元'
table_2026.rows[2].cells[2].text = '专区成本'

table_2026.rows[3].cells[0].text = '平均成本'
table_2026.rows[3].cells[1].text = str(int(df_2026['专区成本'].mean())) + '元'
table_2026.rows[3].cells[2].text = '无超标'

table_2026.rows[4].cells[0].text = '最高成本'
table_2026.rows[4].cells[1].text = str(int(df_2026['专区成本'].max())) + '元'
table_2026.rows[4].cells[2].text = '低于32,000基线'

table_2026.rows[5].cells[0].text = '最低成本'
table_2026.rows[5].cells[1].text = str(int(df_2026['专区成本'].min())) + '元'
table_2026.rows[5].cells[2].text = '-'

# 3.2 26年度专区明细
doc.add_heading('3.2 26年度新开设专区明细', 2)
table_detail = doc.add_table(rows=len(df_2026)+1, cols=5)
table_detail.style = 'Light Grid Accent 1'
headers_detail = ['合同编号', '专区名称', '省份', '接入时间', '专区成本']
for i, header in enumerate(headers_detail):
    table_detail.rows[0].cells[i].text = header
    table_detail.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

df_2026_sorted = df_2026.sort_values('专区成本', ascending=False)
for idx, (_, row) in enumerate(df_2026_sorted.iterrows(), 1):
    table_detail.rows[idx].cells[0].text = str(row['合同编号'])
    table_detail.rows[idx].cells[1].text = str(row[' 原专区名称'])[:18]
    table_detail.rows[idx].cells[2].text = str(row['所属省份'])
    table_detail.rows[idx].cells[3].text = str(row['确认接入时间'])[:10]
    table_detail.rows[idx].cells[4].text = str(int(row['专区成本']))

doc.add_page_break()

# 四、重点工作规划
doc.add_heading('四、重点工作规划', 1)

doc.add_heading('4.1 成本管控问题专区分析', 2)
doc.add_paragraph('责任人：XXX    完成时间：2026年4月30日')
doc.add_paragraph('问题专区：成本TOP10中的超标专区')
doc.add_paragraph('成本复盘计划：')
doc.add_paragraph('分析成本构成：人力成本、系统对接成本、运维成本占比', style='List Bullet')
doc.add_par