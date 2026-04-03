#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修改原有报告文档
添加指标七，修正数据重复问题
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 打开原有文档
doc_path = r'D:\work\运营中心\yxzx\新点e交易相关材料\日常数据运维工具\temp\e交易专区收益成本统计报告.docx'
doc = Document(doc_path)

# 设置中文字体函数
def set_chinese_font(run):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 在"重要说明"段落后插入新内容
# 找到"二、核心数据概览"章节，在其前面插入指标七的说明

# 遍历所有段落，找到插入位置
insert_index = None
for i, para in enumerate(doc.paragraphs):
    if '二、核心数据概览' in para.text:
        insert_index = i
        break

if insert_index:
    # 在"二、核心数据概览"前插入新内容
    
    # 1. 添加指标七说明
    p_new = doc.paragraphs[insert_index]._element
    p_new.getparent().insert(p_new.getparent().index(p_new), 
        doc.add_heading('一、指标说明（补充）', level=1)._element)
    
    # 添加指标七说明段落
    p7 = doc.add_paragraph()
    p7_run = p7.add_run('【指标七（新增）】接入超过1年，总收益为0的专区')
    p7_run.bold = True
    p7_run.font.color.rgb = RGBColor(192, 0, 0)
    p7_run.font.size = Pt(12)
    set_chinese_font(p7_run)
    
    p7_text = p7.add_run('\n\n经过深入分析，我们发现部分专区虽然接入时间超过1年，但总收益为0，这类专区需要特别关注。指标七专门统计这类专区，共识别出19个零收益专区。')
    p7_text.font.size = Pt(11)
    set_chinese_font(p7_text)
    
    # 2. 添加数据包含关系说明
    p_relation = doc.add_paragraph()
    p_relation_run = p_relation.add_run('【数据包含关系说明】')
    p_relation_run.bold = True
    p_relation_run.font.color.rgb = RGBColor(192, 0, 0)
    p_relation_run.font.size = Pt(12)
    set_chinese_font(p_relation_run)
    
    relation_text = '''\n\n为避免重复统计，本报告明确以下包含关系：
• 指标三（总收益<5w）是指标一（总收益<10w）的子集
• 指标七（总收益=0）是指标一（总收益<10w）的子集
• 指标四（25年收益<5w）是指标二（25年收益<10w）的子集

因此，指标一 = 指标一独有（5-10w）+ 指标三（<5w）+ 指标七（=0）
经过去重处理，实际需要关注的专区总数为83个。'''
    
    p_relation_text = p_relation.add_run(relation_text)
    p_relation_text.font.size = Pt(11)
    set_chinese_font(p_relation_text)

# 找到"2.1 六大指标统计"，修改为"2.1 七大指标统计"并添加指标七
for para in doc.paragraphs:
    if '2.1 六大指标统计' in para.text:
        para.text = '2.1 七大指标统计（原始数据）'
        for run in para.runs:
            set_chinese_font(run)
        break

# 找到第一个表格（原始指标表格），在最后一行添加指标七
for table in doc.tables:
    # 检查是否是原始指标表格（有5列）
    if len(table.columns) == 5:
        # 在最后一行后添加指标七
        new_row = table.add_row()
        cells_data = ['指标七（新增）', '超1年且总收益为0', '19', '属于指标一', '新增']
        for i, text in enumerate(cells_data):
            cell = new_row.cells[i]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    set_chinese_font(run)
                if i in [0, 2, 3, 4]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        break

# 找到"2.2 去重后的风险等级分布"，添加黑色等级
for i, para in enumerate(doc.paragraphs):
    if '2.2 去重后的风险等级分布' in para.text:
        # 在其后的表格中添加黑色等级行
        # 找到下一个表格
        for j in range(i+1, len(doc.paragraphs)):
            # 查找表格
            if j < len(doc.paragraphs):
                # 尝试找到表格
                pass
        break

# 找到第二个表格（去重后表格），添加黑色等级
for table in doc.tables:
    # 检查是否是去重后表格（有4列）
    if len(table.columns) == 4 and len(table.rows) >= 5:
        # 在"灰色"行后插入"黑色"行
        # 找到"灰色-流失风险"行
        for i, row in enumerate(table.rows):
            if '灰色' in row.cells[0].text:
                # 在灰色行后添加黑色行
                # 由于docx库不支持直接插入行，我们在表格后添加说明
                pass
        
        # 修改"合计"行的说明
        for row in table.rows:
            if '合计' in row.cells[0].text:
                row.cells[3].text = '去重后总数（指标七已包含在红色中）'
                for paragraph in row.cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        set_chinese_font(run)
                break
        break

# 在"去重后的风险等级分布"表格后添加黑色等级说明
for i, para in enumerate(doc.paragraphs):
    if '2.2 去重后的风险等级分布' in para.text:
        # 找到表格后的段落
        table_found = False
        for j in range(i+1, min(i+20, len(doc.paragraphs))):
            # 跳过表格
            if not table_found:
                # 检查是否是表格
                pass
        break

# 添加黑色等级说明段落（在所有表格后）
black_para = doc.add_paragraph()
black_run = black_para.add_run('【补充说明】')
black_run.bold = True
black_run.font.color.rgb = RGBColor(0, 0, 0)
black_run.font.size = Pt(12)
set_chinese_font(black_run)

black_text = black_para.add_run('\n\n⚫ 黑色-零收益（超1年且总收益为0）：19个专区\n')
black_text.font.size = Pt(11)
set_chinese_font(black_text)

black_desc = black_para.add_run('说明：这19个专区属于红色等级的子集，是长期运营但未产生任何收益的专区，需要重点排查原因，考虑是否下线或重新激活。')
black_desc.font.size = Pt(11)
set_chinese_font(black_desc)

# 修改"四、后续处理建议"中的红色等级处理数量
for para in doc.paragraphs:
    if '4.1 红色等级专区' in para.text:
        # 修改数量
        para.text = '4.1 红色等级专区（66个，其中黑色19个）'
        for run in para.runs:
            set_chinese_font(run)
        break

# 在红色等级处理措施中添加对黑色专区的特别说明
for i, para in enumerate(doc.paragraphs):
    if '4.1 红色等级专区' in para.text:
        # 找到处理措施段落
        for j in range(i+1, min(i+10, len(doc.paragraphs))):
            if '逐一排查专区运营状态' in doc.paragraphs[j].text:
                # 在其前插入黑色专区特别说明
                special_para = doc.paragraphs[j]._element
                special_p = doc.add_paragraph()
                special_run = special_p.add_run('【黑色专区（19个）特别处理】')
                special_run.bold = True
                special_run.font.color.rgb = RGBColor(0, 0, 0)
                set_chinese_font(special_run)
                
                special_text = special_p.add_run('\n对于总收益为0的19个黑色专区，建议：\n• 立即排查是否仍在正常运营\n• 分析零收益原因：是业务未开展、客户未入驻还是系统问题\n• 评估是否值得继续投入运营成本\n• 对于无挽救价值的专区，果断下线释放资源')
                set_chinese_font(special_text)
                
                special_p._element.getparent().insert(
                    special_p._element.getparent().index(special_para),
                    special_p._element
                )
                break
        break

# 保存修改后的文档
output_path = r'D:\work\运营中心\yxzx\新点e交易相关材料\日常数据运维工具\temp\e交易专区收益成本统计报告_修正版.docx'
doc.save(output_path)
print(f'文档已修改并保存: {output_path}')
print('\n修改内容：')
print('1. 添加指标七：接入超过1年，总收益为0（19个专区）')
print('2. 添加数据包含关系说明')
print('3. 添加黑色-零收益风险等级说明')
print('4. 修正红色等级处理建议，添加黑色专区特别处理')
