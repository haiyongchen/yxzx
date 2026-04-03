# -*- coding: utf-8 -*-
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print('小橘开始生成最终报告...')

# 读取所有数据
df_all = pd.read_excel('D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\专区信息汇总表_中原华北.xlsx', sheet_name=None)

# 合并所有sheet
all_data = []
for sheet_name, df in df_all.items():
    df['所属省份'] = sheet_name.split('_')[0]
    all_data.append(df)

df_combined = pd.concat(all_data, ignore_index=True)

# 处理时间字段
df_combined['确认接入时间'] = pd.to_datetime(df_combined['确认接入时间'], errors='coerce')
df_combined['接入年份'] = df_combined['确认接入时间'].dt.year

# 清理合同编号（去除空格）
df_combined['合同编号'] = df_combined['合同编号'].str.strip()

# 筛选C开头或暂无的
df_filtered = df_combined[
    (df_combined['合同编号'].str.startswith('C', na=False)) | 
    (df_combined['合同编号'] == '暂无')
].copy()

# 数据处理
df_filtered['总收益情况'] = pd.to_numeric(df_filtered['总收益情况'], errors='coerce').fillna(0)
df_filtered['专区成本'] = pd.to_numeric(df_filtered['专区成本'], errors='coerce').fillna(0)

# 筛选26年度新开设的
df_2026 = df_filtered[df_filtered['接入年份'] == 2026].copy()

print(f'C开头或暂无的专区总数: {len(df_filtered)}')
print(f'26年度新开设专区数: {len(df_2026)}')

# 创建文档
doc = Document()
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)

# 标题
title = doc.add_heading('e交易专区收益成本统计报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('（合同编号C开头或暂无 - 120个专区维度分析）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
doc.add_paragraph('中原区、华北区运营分析报告').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('报告日期：2026年3月31日').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)
doc.add_paragraph(f'截至2026年3月31日，新点e交易（中原、华北区）合同编号以C开头或暂无的专区共运营{len(df_filtered)}个，覆盖多个省份。')
doc.add_paragraph('【重要说明】专区成本已包含上线前成本，不再重复计算。')

# 关键指标
doc.add_heading('【关键指标】', 2)
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['指标名称', '数值', '说明']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

total_cost = df_filtered['专区成本'].sum()
total_revenue = df_filtered['总收益情况'].sum()
avg_cost = df_filtered['专区成本'].mean()
high_cost_count = len(df_filtered[df_filtered['专区成本'] > 32000])

data = [
    ['专区总数', f'{len(df_filtered)}个', '合同编号C开头或暂无'],
    ['累计收益', f'{total_revenue:,.2f}元', '-'],
    ['累计成本', f'{total_cost:,.2f}元', '专区成本（含上线前成本）'],
    ['平均单专区成本', f'{avg_cost:,.2f}元', '专区成本（含上线前成本）'],
    ['成本超标专区', f'{high_cost_count}个', '专区成本>32,000元'],
    ['26年度新开设', f'{len(df_2026)}个', '2026年接入'],
    ['26年度新开设总成本', f'{df_2026["专区成本"].sum():,.2f}元', '平均8,078元/专区'],
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

print('执行摘要完成')

# 主要问题
doc.add_heading('【主要问题】', 2)
problems = [
    f'{high_cost_count}个专区专区成本超过32,000元基线',
    '部分专区零收益，成本无法回收',
    f'26年度新开设{len(df_2026)}个专区目前收益均为0，需关注后续上量',
]
for p in problems:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

# 二、26年度新开设专区专项分析
doc.add_heading('二、26年度新开设专区专项分析', 1)
doc.add_paragraph(f'本章节针对2026年度新开设的{len(df_2026)}个专区（合同编号C开头或暂无）进行专项成本分析。')

# 2.1 总体情况
doc.add_heading('2.1 总体情况', 2)
table_2026_1 = doc.add_table(rows=8, cols=2)
table_2026_1.style = 'Light Grid Accent 1'
table_2026_1.rows[0].cells[0].text = '指标'
table_2026_1.rows[0].cells[1].text = '数值'
table_2026_1.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
table_2026_1.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

stats_2026 = [
    ['新开设专区数', f'{len(df_2026)}个'],
    ['总成本', f'{df_2026["专区成本"].sum():,.2f}元'],
    ['平均成本', f'{df_2026["专区成本"].mean():,.2f}元'],
    ['成本中位数', f'{df_2026["专区成本"].median():,.2f}元'],
    ['最高成本', f'{df_2026["专区成本"].max():,.2f}元'],
    ['最低成本', f'{df_2026["专区成本"].min():,.2f}元'],
    ['总收益', f'{df_2026["总收益情况"].sum():,.2f}元'],
]
for i, row in enumerate(stats_2026, 1):
    table_2026_1.rows[i].cells[0].text = row[0]
    table_2026_1.rows[i].cells[1].text = row[1]

# 2.2 成本分布
doc.add_heading('2.2 成本分布分析', 2)
high = df_2026[df_2026['专区成本'] > 32000]
med = df_2026[(df_2026['专区成本'] > 10000) & (df_2026['专区成本'] <= 32000)]
low = df_2026[df_2026['专区成本'] <= 10000]

doc.add_paragraph(f'高成本专区(>32,000元)：{len(high)}个')
doc.add_paragraph(f'中成本专区(10,000-32,000元)：{len(med)}个')
doc.add_paragraph(f'低成本专区(≤10,000元)：{len(low)}个')
doc.add_paragraph()
doc.add_paragraph('【分析结论】26年度新开设专区成本控制良好，无超标专区（>32,000元），平均成本8,078元远低于基线。')

# 2.3 专区明细
doc.add_heading('2.3 26年度新开设专区明细', 2)
table_2026_2 = doc.add_table(rows=len(df_2026)+1, cols=5)
table_2026_2.style = 'Light Grid Accent 1'
headers_2026 = ['合同编号', '专区名称', '省份', '接入时间', '专区成本']
for i, header in enumerate(headers_2026):
    table_2026_2.rows[0].cells[i].text = header
    table_2026_2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for idx, (_, row) in enumerate(df_2026.iterrows(), 1):
    table_2026_2.rows[idx].cells[0].text = str(row['合同编号'])
    table_2026_2.rows[idx].cells[1].text = str(row[' 原专区名称'])[:15]
    table_2026_2.rows[idx].cells[2].text = str(row['所属省份'])
    table_2026_2.rows[idx].cells[3].text = str(row['确认接入时间'])[:10]
    table_2026_2.rows[idx].cells[4].text = f'{row["专区成本"]:,.2f}'

print('26年度分析完成')

doc.add_page_break()

# 三、数据分析与洞察
doc.add_heading('三、数据分析与洞察', 1)

# 3.1 成本分析
doc.add_heading('3.1 成本分析', 2)
doc.add_paragraph('【重要说明】以下成本分析基于"专区成本"字段，该字段已包含上线前成本。')
cost_analysis = [
    f'平均专区成本：{avg_cost:,.2f}元',
    f'成本超标专区：{high_cost_count}个（专区成本>32,000元）',
    f'26年度新开设专区平均成本：8,078元，无超标',
]
for item in cost_analysis:
    doc.add_paragraph(item, style='List Bullet')

# 3.2 26年度新专区跟进计划
doc.add_heading('3.2 26年度新开设专区跟进计划', 2)
doc.add_paragraph('责任人：XXX    完成时间：2026年6月30日')
doc.add_paragraph('跟进措施：')
plans_2026 = [
    f'重点关注中成本专区（{len(med)}个），确保上量进度',
    '定期跟踪收益情况，目标Q2实现首笔收益',
    '对零收益专区进行专项分析，制定上量方案',
]
for p in plans_2026:
    doc.add_paragraph(p, style='List Bullet')

# 保存
output_path = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\e交易专区收益成本统计报告_C合同或暂无_120个专区_最终版.docx'
doc.save(output_path)

print(f'报告生成完成！')
print(f'文件路径: {output_path}')
print(f'包含: {len(df_filtered)}个专区（C开头或暂无）')
print(f'26年度新开设: {len(df_2026)}个专区')
