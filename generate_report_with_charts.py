#!/usr/bin/env python3
"""
e交易专区收益成本统计报告生成器（增强版）
支持图表生成和新增指标
"""

import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # 使用非交互式后端
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os
import sys

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def analyze_data(file_path):
    """分析数据并返回统计结果"""
    df = pd.read_excel(file_path)
    df['确认接入时间'] = pd.to_datetime(df['确认接入时间'], errors='coerce')
    
    current_date = pd.to_datetime('2026-04-08')
    half_year_ago = current_date - timedelta(days=180)
    
    stats = {
        'total': len(df),
        'indicator1': len(df[(df['确认接入时间'] < '2025-04-08') & (df['总收益情况'] == 0)]),
        'indicator2': len(df[(df['确认接入时间'] < '2025-04-08') & (df['总收益情况'] > 0) & (df['总收益情况'] < 100000)]),
        'indicator3': len(df[(df['确认接入时间'] < '2025-04-08') & (df['总收益情况'] > 0) & (df['总收益情况'] < 50000)]),
        'indicator4': len(df[(df['确认接入时间'] < '2025-01-01') & (df['25年收益情况'] > 0) & (df['25年收益情况'] < 100000)]),
        'indicator5': len(df[(df['确认接入时间'] < '2025-01-01') & (df['25年收益情况'] > 0) & (df['25年收益情况'] < 50000)]),
        'indicator6': len(df[df['26年收益情况'] > 0]),
        'indicator7': len(df[(df['25年收益情况'] > 0) & (df['26年收益情况'] == 0)]),
        'indicator8': len(df[(df['确认接入时间'] < half_year_ago) & (df['总收益情况'] == 0)]),  # 新增指标
    }
    
    # 风险等级分布数据（用于饼图）
    risk_data = {
        '红色-重点关注': stats['indicator1'] + stats['indicator2'],
        '橙色-需改进': stats['indicator4'],
        '黄色-观察': stats['indicator6'],
        '灰色-流失风险': stats['indicator7'],
    }
    
    # 25年收益分布数据（用于柱状图）
    df_25 = df[(df['确认接入时间'] < '2025-01-01') & (df['25年收益情况'] > 0) & (df['25年收益情况'] < 100000)]
    revenue_bins = [0, 30000, 50000, 70000, 100000]
    revenue_labels = ['0-3万', '3-5万', '5-7万', '7-10万']
    df_25['收益区间'] = pd.cut(df_25['25年收益情况'], bins=revenue_bins, labels=revenue_labels)
    revenue_dist = df_25['收益区间'].value_counts().sort_index()
    
    # TOP10流失风险专区
    df_loss = df[(df['25年收益情况'] > 0) & (df['26年收益情况'] == 0)].copy()
    df_loss = df_loss.sort_values('25年收益情况', ascending=False).head(10)
    
    return stats, risk_data, revenue_dist, df_loss, df

def create_chart1_risk_distribution(risk_data, output_path):
    """生成风险等级分布饼图"""
    colors = ['#FF6B6B', '#FFA500', '#FFD700', '#808080']
    
    fig, ax = plt.subplots(figsize=(10, 6))
    wedges, texts, autotexts = ax.pie(
        risk_data.values(), 
        labels=risk_data.keys(),
        autopct='%1.1f%%',
        colors=colors,
        startangle=90,
        textprops={'fontsize': 12}
    )
    
    # 设置百分比文字为白色
    for autotext in autotexts:
        autotext.set_color('white')
        autotext.set_fontweight('bold')
    
    ax.set_title('e交易专区风险等级分布', fontsize=16, fontweight='bold', pad=20)
    
    # 添加图例
    ax.legend(wedges, [f'{k}: {v}个' for k, v in risk_data.items()],
              title="风险等级",
              loc="center left",
              bbox_to_anchor=(1, 0, 0.5, 1))
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'图表1已生成: {output_path}')

def create_chart2_revenue_distribution(revenue_dist, output_path):
    """生成25年收益分布柱状图"""
    fig, ax = plt.subplots(figsize=(10, 6))
    
    bars = ax.bar(revenue_dist.index, revenue_dist.values, color='#4A90D9', edgecolor='black')
    
    # 在柱子上添加数值标签
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{int(height)}',
                ha='center', va='bottom', fontsize=11)
    
    ax.set_xlabel('25年收益区间', fontsize=12)
    ax.set_ylabel('专区数量', fontsize=12)
    ax.set_title('25年前接入专区收益分布（25年总收益<10万）', fontsize=14, fontweight='bold')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'图表2已生成: {output_path}')

def create_chart3_top10_loss(df_loss, output_path):
    """生成TOP10流失风险专区柱状图"""
    fig, ax = plt.subplots(figsize=(12, 6))
    
    # 水平柱状图
    y_pos = range(len(df_loss))
    bars = ax.barh(y_pos, df_loss['25年收益情况'].values, color='#E74C3C', edgecolor='black')
    
    # 设置y轴标签
    ax.set_yticks(y_pos)
    ax.set_yticklabels(df_loss['专区名称'].values, fontsize=10)
    
    # 在柱子上添加数值标签
    for i, bar in enumerate(bars):
        width = bar.get_width()
        ax.text(width, bar.get_y() + bar.get_height()/2.,
                f' {int(width):,}',
                ha='left', va='center', fontsize=9)
    
    ax.set_xlabel('25年收益（元）', fontsize=12)
    ax.set_title('TOP10 流失风险专区（25年有收益但26年无）', fontsize=14, fontweight='bold')
    ax.grid(axis='x', alpha=0.3, linestyle='--')
    
    # 反转y轴，使最大值在顶部
    ax.invert_yaxis()
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight')
    plt.close()
    print(f'图表3已生成: {output_path}')

def create_word_report(stats, risk_data, revenue_dist, df_loss, df, output_path, chart_paths):
    """生成Word报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('新点e交易专区（中原/华北）收益成本', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 概述段落
    doc.add_paragraph(
        f'近期基于新点e交易平台各专区的运营数据，从收益、成本、时间维度进行深度分析。'
        f'通过部分指标的统计分析，我们发现当前平台存在一定数量的零收益、低收益专区，需要引起重点关注并采取针对性措施。'
    )
    
    doc.add_paragraph(
        f'数据统计范围涵盖中原、华北等区域共{stats["total"]}个专区，重点分析确认接入时间、25年总收益、26年总收益、总成本等关键指标。'
        f'通过建立风险分级体系（红色-重点关注、橙色-需改进、黄色-观察、灰色-流失风险）。'
    )
    
    doc.add_paragraph(
        f'本次统计识别出一批需重点关注的低收益专区，其中纳入红色等级预警的有 {stats["indicator1"] + stats["indicator2"]} 个、'
        f'橙色等级预警的有 {stats["indicator4"]} 个（两类等级存在部分专区重叠），另有灰色等级 {stats["indicator7"]} 个。'
        f'从数据趋势来看，部分早期接入的专区已出现运营效率低下、客户活跃度下降等问题。'
    )
    
    doc.add_paragraph(
        f'对于长期无收益且无可挽救价值的专区，建议果断下线以节约运营成本；对于有潜力的专区，加大资源投入和运营支持。'
    )
    
    # 一、核心数据概览
    doc.add_heading('一、核心数据概览', level=1)
    doc.add_heading('1.1 七大指标统计', level=2)
    
    doc.add_paragraph(
        f'从上表可以看出，接入超1年，但长期零收益专区为0的有{stats["indicator1"]}个，'
        f'长期低收益专区（总收益10w以下）共计{stats["indicator2"]}个，占比较高，'
        f'因收益10w以下的专区过多，故进一步分析收益5w以下的专区，数量为{stats["indicator3"]}个；'
        f'25年前接入但收益不达标的专区（25年总收益在10w以下）共计{stats["indicator4"]}个，'
        f'其中25年总收益在5w以下的有{stats["indicator5"]}个，反映出早期接入专区的运营效率问题；'
        f'特别值得关注的是，有{stats["indicator7"]}个专区在25年产生收益但26年暂无收益，存在客户流失风险；'
        f'此外，接入超过半年但仍无收益的专区有{stats["indicator8"]}个，需要重点关注。'
    )
    
    # 指标统计表格（七大指标）
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['指标', '条件', '数量', '风险等级', '备注']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # 数据行
    data_rows = [
        ['指标一', '接入超过1年，总收益为0', str(stats['indicator1']), '红色', '重点关注'],
        ['指标二', '接入超过1年，0<总收益<10w', str(stats['indicator2']), '红色', '重点关注'],
        ['指标三', '接入超过1年，0<总收益<5w', str(stats['indicator3']), '红色', '重点关注'],
        ['指标四', '25年前接入，0<25年收益<10w', str(stats['indicator4']), '橙色', '需改进'],
        ['指标五', '25年前接入，0<25年收益<5w', str(stats['indicator5']), '橙色', '需改进'],
        ['指标六', '26年产生收益', str(stats['indicator6']), '黄色', '观察'],
        ['指标七', '25年有收益，26年无', str(stats['indicator7']), '灰色', '流失风险'],
        ['指标八', '接入超半年，总收益为0', str(stats['indicator8']), '深红', '紧急关注'],
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        for j, cell_data in enumerate(row_data):
