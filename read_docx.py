import os
import sys

sys.stdout.reconfigure(encoding='utf-8')

base_dir = r'D:\work\运营中心\yxzx\阳光优采\需求\在线支付'
docx_path = None
for f in os.listdir(base_dir):
    if f.endswith('.docx'):
        docx_path = os.path.join(base_dir, f)
        break

print(f"Reading: {docx_path}\n")

from docx import Document
doc = Document(docx_path)

# Read all paragraphs
print("=== 段落内容 ===")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else 'Normal'
        print(f"[{style}] {para.text}")

# Read all tables
print("\n=== 表格内容 ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {t_idx + 1} ---")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            print(f"  行{r_idx + 1}: {' | '.join(cells)}")
