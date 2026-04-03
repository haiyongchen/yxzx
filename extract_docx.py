# -*- coding: utf-8 -*-
from docx import Document
import zipfile
import os

doc_path = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\e交易专区收益成本统计报告.docx'
output_path = 'D:\\openclaw-workspace\\report_content.txt'

# 读取Word文档
doc = Document(doc_path)

with open(output_path, 'w', encoding='utf-8') as f:
    f.write("=" * 80 + "\n")
    f.write("e交易专区收益成本统计报告 - 内容分析\n")
    f.write("=" * 80 + "\n")
    
    # 提取所有段落
    f.write("\n【文档段落内容】\n\n")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            f.write(f"段落 {i+1}: {para.text}\n")
    
    # 提取所有表格
    f.write("\n\n【文档表格内容】\n\n")
    for table_idx, table in enumerate(doc.tables):
        f.write(f"\n--- 表格 {table_idx + 1} ---\n")
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text for cell in row.cells]
            f.write(f"行 {row_idx + 1}: {row_data}\n")

# 提取嵌入的Excel文件
extract_path = 'D:\\openclaw-workspace\\extracted_files'
os.makedirs(extract_path, exist_ok=True)

with zipfile.ZipFile(doc_path, 'r') as zip_ref:
    # 查找嵌入的文件
    for file in zip_ref.namelist():
        if 'embeddings' in file:
            zip_ref.extract(file, extract_path)
            print(f"提取文件: {file}")

print(f"\n文档内容已保存到: {output_path}")
