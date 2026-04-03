#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
e交易专区收益成本统计报告 - 修正版
数据关系：指标三⊆指标一，指标四⊆指标二
去重后统计：红色66个，橙色6个，灰色4个，黄色7个，合计83个
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(run):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 添加标题
title = doc.add_heading('e交易专区收益成本统计报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(31, 78, 121)
    set_chinese_font(run)

# 添加日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run('报告日期：2026年4月2日（修正版）')
date_run.font.size = Pt(12)
date_run.font.color.rgb = RGBColor(102, 102, 102)
set_chinese_font(date_run)

doc.add_paragraph()  # 空行

# 一、报告概述
doc.add_heading('一、报告概述', level=1)

p1 = doc.add_paragraph()
p1_text = '''本报告基于e交易平台各专区的运营数据，从收益、成本、时间维度进行深度分析，旨在识别低效运营专区，为管理层决策提供数据支撑。

数据统计范围涵盖中原、华北等区域共147个专区。通过对六大核心指标的统计分析，我们发现指标之间存在包含关系，需要进行去重处理以获得真实的风险分布情况。'''
p1_run = p1.add_run(p1_text)
p1_run.font.size = Pt(11)
set_chinese_font(p1_run)

# 重要说明
important = doc.add_paragraph()
important_run = important.add_run('【重要说明】')
important_run.bold = True
important_run.font.color.rgb = RGBColor(192, 0, 0)
important_run.font.size = Pt(11)
set_chinese_font(important_run)

important_text = important.add_run(' 本报告中的指标存在包含关系：指标三（总收益<5w）是指标一（总收益<10w）的子集，指标四（25年收益<5w）是指标二（25年收益<10w）的子集。经过去重处理，实际需要关注的专区总数为83个。')
important_text.font.size = Pt(11)
set_chinese_font(important_text)

doc.add_paragraph()

# 二、核心数据概览
doc.add_heading('二、核心数据概览', level=1)
doc.add_heading('2.1 六大指标统计（原始数据）', level=2)

# 原始数据表格
table1 = doc.add_table(rows=7, cols=5)
table1.style = 'Table Grid'

# 表头
headers = ['指标', '条件', '数量', '包含关系', '备注']
for i, header in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
            set_chinese_font(run)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置背景色
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '1F4E79')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# 数据行
data_rows = [
    ['指标一', '超1年且总收益<10w', '66', '包含指标三', ''],
    ['指标二', '25年前且25年收益<10w', '64', '包含指标四', ''],
    ['指标三', '超1年且总收益<5w', '56', '属于指标一', '子集'],
    ['指标四', '25年前且25年收益<5w', '60', '属于指标二', '子集'],
    ['指标五', '26年产生收益', '33', '-', ''],
    ['指标六', '25年有收益但26年无', '20', '-', ''],
]

for i, row_data in enumerate(data_rows, 1):
    for j, text in enumerate(row_data):
        cell = table1.rows[i].cells[j]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                set_chinese_font(run)
            if j in [0, 2, 3, 4]:  # 居中列
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 说明文字
explain = doc.add_paragraph()
explain_text = explain.add_run('从上表可以看出，指标之间存在明显的包含关系：指标三（56个）全部包含在指标一（66个）中，指标四（60个）全部包含在指标二（64个）中。这意味着如果简单相加会导致重复统计，需要进行去重处理。')
explain_text.font.size = Pt(11)
set_chinese_font(explain_text)

doc.add_paragraph()

# 2.2 去重后的风险等级分布
doc.add_heading('2.2 去重后的风险等级分布', level=2)

# 去重后表格
table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Table Grid'

# 表头
headers2 = ['风险等级', '条件', '数量', '说明']
for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_chinese_font(run)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '1F4E79')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# 数据行 - 去重后
deduped_rows = [
    ['🔴 红色-重点关注', '超1年且总收益<10w', '66', '指标一全部'],
    ['🟠 橙色-需改进', '25年前且25年收益<10w（排除红色）', '6', '指标二独有'],
    ['⚪ 灰色-流失风险', '25年有但26年无（排除红橙）', '4', '指标六排除重叠'],
    ['🟡 黄色-观察', '26年有收益（排除红橙灰）', '7', '指标五排除重叠'],
    ['合计', '不重复专区总数', '83', '去重后总数'],
]

colors = ['C00000', 'ED7D31', '7F7F7F', 'FFC000', '000000']

for i, (row_data, color) in enumerate(zip(deduped_rows, colors), 1):
    for j, text in enumerate(row_data):
        cell = table2.rows[i].cells[j]
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)
                if j == 0:  # 风险等级列
                    run.bold = True
                    run.font.color.rgb = RGBColor(int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16))
                if i == 5:  # 合计行
                    run.bold = True
                set_chinese_font(run)
            if j in [0, 2, 3]:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 说明文字
explain2 = doc.add_paragraph()
explain2_text = explain2.add_run('经过去重处理，实际需要关注的专区总数为83个。其中红色等级（重点关注）66个，占比79.5%；橙色等级（需改进）6个；灰色等级（流失风险）4个；黄色等级（观察）7个。这一分布特征表明，平台存在大量长期低收益专区，需要系统性整改。')
explain2_text.font.size = Pt(11)
set_chinese_font(explain2_text)

doc.add_paragraph()

# 插入图表
doc.add_heading('2.3 风险等级分布图', level=2)

try:
    doc.add_picture('D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\chart1_risk_distribution.png', width=Inches(5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph('图1：e交易专区风险等级分布（去重后）')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.size = Pt(10)
        run.italic = True
        set_chinese_font(run)
except:
    pass

doc.add_paragraph()

# 三、详细分类说明
doc.add_heading('三、详细分类说明', level=1)

p_detail = doc.add_paragraph()
p_detail_text = '''基于数据包含关系的分析，我们将83个需要关注的专区进行如下分类：

1. 【超1年且总收益5-10w】10个（指标一独有）
   这些专区接入时间较长，收益处于中等偏低水平，有提升空间。

2. 【超1年且总收益<5w】56个（指标三）
   长期低收益专区，需要重点分析原因，考虑优化或下线。

3. 【25年前且25年收益5-10w】4个（指标二独有）
   25年前接入，25年收益处于中等水平，需关注26年表现。

4. 【25年前且25年收益<5w】60个（指标四）
   早期接入但持续低收益，需要评估是否继续运营。

5. 【26年产生收益】7个（指标五排除重叠后）
   当前运营状态良好，继续保持观察。

6. 【25年有收益但26年无】4个（指标六排除重叠后）
   存在流失风险，需要紧急跟进。'''

p_detail_run = p_detail.add_run(p_detail_text)
p_detail_run.font.size = Pt(11)
set_chinese_font(p_detail_run)

doc.add_paragraph()

# 四、后续处理建议
doc.add_heading('四、后续处理建议', level=1)

doc.add_heading('4.1 红色等级专区（66个）', level=2)
p_red = doc.add_paragraph()
p_red_text = p_red.add_run('处理措施：')
p_red_text.bold = True
p_red_text.font.size = Pt(11)
set_chinese_font(p_red_text)

red_measures = '''
• 逐一排查专区运营状态，确认是否仍在正常服务客户
• 对于长期无交易的专区，评估是否继续投入运营成本
• 制定专区下线或合并方案，释放服务器和运维资源
• 对于仍有潜力的专区，制定专项提升计划'''

p_red2 = doc.add_paragraph(red_measures)
for run in p_red2.runs:
    run.font.size = Pt(11)
    set_chinese_font(run)

doc.add_heading('4.2 灰色等级专区（4个）', level=2)
p_grey = doc.add_paragraph()
p_grey_text = p_grey.add_run('处理措施：')
p_grey_text.bold = True
p_grey_text.font.size = Pt(11)
set_chinese_font(p_grey_text)

grey_measures = '''
• 优先安排客户经理进行电话或实地回访
• 了解客户流失原因：是业务暂停、平台迁移还是服务问题
• 针对可挽回客户，制定专属优惠政策或服务升级方案
• 建立客户流失预警机制，定期监控专区活跃度'''

p_grey2 = doc.add_paragraph(grey_measures)
for run in p_grey2.runs:
    run.font.size = Pt(11)
    set_chinese_font(run)

doc.add_heading('4.3 橙色等级专区（6个）', level=2)
p_orange = doc.add_paragraph()
p_orange_text = p_orange.add_run('处理措施：')
p_orange_text.bold = True
p_orange_text.font.size = Pt(11)
set_chinese_font(p_orange_text)

orange_measures = '''
• 分析专区所在行业特点和区域市场环境
• 对比同类高收益专区的运营模式
• 制定冲刺方案争取突破10万门槛
• 定期组织运营培训，提升运营能力'''

p_orange2 = doc.add_paragraph(orange_measures)
for run in p_orange2.runs:
    run.font.size = Pt(11)
    set_chinese_font(run)

# 五、总结
doc.add_heading('五、总结', level=1)

p_summary = doc.add_paragraph()
p_summary_text = '''本次统计经过去重处理后，实际需要关注的专区总数为83个，而非简单相加的306个。其中红色等级66个（占比79.5%），是最需要重点关注的群体。

建议成立专项工作组，按照"先急后缓、先大后小"的原则，分批推进低收益专区的整改提升工作。对于长期无收益且无可挽救价值的专区，建议果断下线以节约运营成本；对于有潜力的专区，加大资源投入和运营支持，力争在2026年下半年实现收益显著提升。

后续将建立月度监控机制，持续跟踪各专区收益变化情况，及时发现问题并调整策略，确保e交易平台整体运营质量稳步提升。'''

p_summary_run = p_summary.add_run(p_summary_text)
p_summary_run.font.size = Pt(11)
set_chinese_font(p_summary_run)

doc.add_paragraph()

# 附录
appendix = doc.add_paragraph()
appendix_run = appendix.add_run('【附录】')
appendix_run.bold = True
appendix_run.font.size = Pt(11)
appendix_run.font.color.rgb = RGBColor(102, 102, 102)
set_chinese_font(appendix_run)

appendix_text = appendix.add_run(' 详细数据请参见附件《专区低收益统计结果.xlsx》和《数据分析结果_修正版.xlsx》，包含六个指标的完整数据清单及去重分析。')
appendix_text.font.size = Pt(10)
appendix_text.font.color.rgb = RGBColor(102, 102, 102)
set_chinese_font(appendix_text)

# 保存文档
output_path = 'D:\\work\\运营中心\\yxzx\\新点e交易相关材料\\日常数据运维工具\\temp\\e交易专区收益成本统计报告_修正版.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')